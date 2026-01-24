#!/usr/bin/env python3
"""
Thesis revision management tool for Tri-CertFramework docs.

Usage:
    ./thesis.py init             Initialize thesis source directory structure
    ./thesis.py split            Split existing thesis into source structure
    ./thesis.py build            Build thesis from source and create new revision
    ./thesis.py major            Increment major version (1.x → 2.0)
    ./thesis.py fetch            Sync config with research directory versions
    ./thesis.py list             List all revisions
    ./thesis.py tree             Show source directory structure
    ./thesis.py show <version>   Show details of a specific revision
    ./thesis.py remove <version> Remove a revision
    ./thesis.py word             Build Word document with cover template
    ./thesis.py pdf              Export thesis to PDF format
    ./thesis.py export           Export thesis to both Word and PDF

Requirements:
    Python 3.8+
    For export commands:
        - pandoc (brew install pandoc)
        - mermaid-cli (npm install -g @mermaid-js/mermaid-cli)
        - For PDF: XeLaTeX (brew install --cask mactex-no-gui)
        - python-docx (pip install python-docx) - for Word export
        - docxcompose (pip install docxcompose) - for Word export
        - Pillow (pip install Pillow) - for image compression

Examples:
    ./thesis.py init             # Create initial directory structure
    ./thesis.py split            # Split existing thesis into source
    ./thesis.py build            # Build new version from source
    ./thesis.py fetch            # Sync config with research dir
    ./thesis.py tree             # Show thesis source structure
    ./thesis.py word             # Build Word document with cover
    ./thesis.py pdf              # Export to PDF for final submission
    ./thesis.py export           # Export to both Word and PDF
"""

import argparse
import configparser
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


# =============================================================================
# Configuration
# =============================================================================

SCRIPT_DIR = Path(__file__).parent
DOCS_DIR = SCRIPT_DIR.parent
CONFIG_FILE = SCRIPT_DIR / "thesis.config"

# These are loaded from config
SOURCE_DIR: Path = DOCS_DIR / "thesis-source"
RESEARCH_DIR: Path = DOCS_DIR / "src" / "content" / "docs" / "research"
INDEX_FILE: Path = RESEARCH_DIR / "index.md"
ASTRO_CONFIG: Path = DOCS_DIR / "astro.config.mjs"

THESIS_FILE_PATTERN = re.compile(r"thesis-v(\d+)-(\d+)\.md")
VERSION_PATTERN = re.compile(r"v?(\d+)[.-](\d+)")
SECTION_PATTERN = re.compile(r"^(\d+(?:\.\d+)*)\s+(.+)$")

# Image path patterns for conversion
# thesis-source uses: ![alt](screenshot/component/file.png)
# Output needs: ![alt](../../../assets/screenshot/component/file.png)
IMAGE_PATH_PATTERN = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
SCREENSHOT_DIR = DOCS_DIR / "src" / "assets" / "screenshot"

# Export-related paths
TEMPLATES_DIR = SCRIPT_DIR / "templates"
OUTPUT_DIR = SCRIPT_DIR / "output"
REFERENCE_DOCX = TEMPLATES_DIR / "reference.docx"
COVER_DOCX = TEMPLATES_DIR / "表紙.docx"


# =============================================================================
# Config Management
# =============================================================================

def load_config() -> configparser.ConfigParser:
    """Load configuration from thesis.config."""
    config = configparser.ConfigParser()
    
    if CONFIG_FILE.exists():
        config.read(CONFIG_FILE, encoding="utf-8")
    else:
        # Create default config
        config["version"] = {"major": "1", "minor": "0"}
        config["metadata"] = {
            "title": "修士論文",
            "subtitle": "デジタル文書の真正性検証: ３層認証アーキテクチャの設計と実装",
            "description": "ゼロ知識証明を用いた分散型証明システム Tri-CertFramework",
        }
        config["paths"] = {
            "source_dir": "thesis-source",
            "output_dir": "src/content/docs/research",
        }
        save_config(config)
    
    return config


def save_config(config: configparser.ConfigParser) -> None:
    """Save configuration to thesis.config."""
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        f.write("# Thesis Management Configuration\n")
        f.write("# This file is managed by thesis.py\n\n")
        config.write(f)


def get_current_version(config: configparser.ConfigParser) -> tuple[int, int]:
    """Get current version from config."""
    return (
        config.getint("version", "major"),
        config.getint("version", "minor"),
    )


def set_version(config: configparser.ConfigParser, major: int, minor: int) -> None:
    """Set version in config."""
    config.set("version", "major", str(major))
    config.set("version", "minor", str(minor))
    save_config(config)


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class SourceFile:
    """Represents a source file in the thesis structure."""
    path: Path
    number: str  # e.g., "1.2.3" or ""
    title: str
    level: int   # 0=frontmatter, 1=chapter, 2=section, 3=subsection
    content: str = ""
    
    @property
    def sort_key(self) -> tuple:
        """Key for sorting files in order."""
        if not self.number:
            # Special files sort by name
            special_order = {
                "_frontmatter": (0, 0, 0, 0),
                "_toc": (0, 0, 0, 1),
                "謝辞": (999, 0, 0, 0),
                "参考文献": (999, 0, 0, 1),
                "付録": (999, 0, 0, 2),
            }
            for key, order in special_order.items():
                if key in self.path.stem:
                    return order
            return (998, 0, 0, 0)
        
        parts = [int(p) for p in self.number.split(".")]
        while len(parts) < 4:
            parts.append(0)
        return tuple(parts)


@dataclass
class Revision:
    """Represents a thesis revision."""
    major: int
    minor: int
    file_path: Path
    is_current: bool = False
    title: str = ""
    
    @property
    def version(self) -> str:
        return f"{self.major}.{self.minor}"
    
    @property
    def slug(self) -> str:
        return f"thesis-v{self.major}-{self.minor}"
    
    @property
    def filename(self) -> str:
        return f"{self.slug}.md"
    
    def __lt__(self, other: "Revision") -> bool:
        return (self.major, self.minor) < (other.major, other.minor)


# =============================================================================
# Source Directory Functions
# =============================================================================

def get_source_dir(config: configparser.ConfigParser) -> Path:
    """Get source directory from config."""
    return DOCS_DIR / config.get("paths", "source_dir")


def parse_dir_or_file_name(name: str) -> tuple[str, str]:
    """Parse directory or file name to extract number and title."""
    # Remove .md extension if present
    name = name.replace(".md", "")
    
    # Try to match section pattern (e.g., "1.2.3 Title")
    match = SECTION_PATTERN.match(name)
    if match:
        return match.group(1), match.group(2)
    
    # Try to match chapter pattern (e.g., "第1章 Title")
    chapter_match = re.match(r"第(\d+)章\s+(.+)", name)
    if chapter_match:
        return chapter_match.group(1), chapter_match.group(2)
    
    return "", name


def collect_source_files(source_dir: Path) -> list[SourceFile]:
    """Collect all source files from the thesis directory structure."""
    files = []
    
    if not source_dir.exists():
        return files
    
    def process_path(path: Path, level: int = 0) -> None:
        if path.is_file() and path.suffix == ".md":
            number, title = parse_dir_or_file_name(path.stem)
            content = path.read_text(encoding="utf-8")
            files.append(SourceFile(
                path=path,
                number=number,
                title=title,
                level=level,
                content=content,
            ))
        elif path.is_dir():
            # Process directory itself (for chapter/section headers)
            dir_number, dir_title = parse_dir_or_file_name(path.name)

            # Check for _chapter.md or _section.md
            header_file = path / "_section.md"
            if not header_file.exists():
                header_file = path / "_chapter.md"

            if header_file.exists():
                content = header_file.read_text(encoding="utf-8")
                files.append(SourceFile(
                    path=header_file,
                    number=dir_number,
                    title=dir_title,
                    level=level,
                    content=content,
                ))
            elif dir_number and dir_title:
                # No header file, but directory has number and title
                # Create a placeholder entry for TOC generation
                files.append(SourceFile(
                    path=path,
                    number=dir_number,
                    title=dir_title,
                    level=level,
                    content="",  # No content, just for TOC
                ))

            # Process children
            children = sorted(path.iterdir(), key=lambda p: (
                0 if p.name.startswith("_") else 1,
                p.name,
            ))
            for child in children:
                if child.name.startswith("_"):
                    continue  # Already processed
                if child.name.startswith("."):
                    continue  # Skip hidden
                process_path(child, level + 1)
    
    # Process all items in source directory
    for item in sorted(source_dir.iterdir()):
        if item.name.startswith("."):
            continue
        # Frontmatter files (starting with _) are level 0
        # Other top-level files (謝辞, 参考文献, etc.) are level 1 (chapter level)
        # Directories start at level 1
        if item.is_file():
            if item.name.startswith("_"):
                process_path(item, 0)  # Frontmatter
            else:
                process_path(item, 1)  # Special chapter-level files
        else:
            process_path(item, 1)
    
    # Sort by section number
    files.sort(key=lambda f: f.sort_key)
    
    return files


def generate_toc(files: list[SourceFile]) -> str:
    """Generate table of contents from source files."""
    toc_lines = ["# 目次\n"]
    
    for f in files:
        # Skip frontmatter and files without titles
        if f.level == 0 or f.path.name.startswith("_") or not f.title:
            continue
        
        # Generate anchor from title
        if f.level == 1 and f.number:
            # Chapter: 第1章 序論
            title = f"第{f.number}章 {f.title}"
            anchor = f"第{f.number}章-{f.title.lower().replace(' ', '-')}"
            toc_lines.append(f"- [{title}](#{anchor})")
        elif f.level == 2 and f.number:
            # Section: 1.1 研究の背景
            title = f"{f.number} {f.title}"
            anchor = f"{f.number.replace('.', '')}-{f.title.lower().replace(' ', '-')}"
            toc_lines.append(f"  - [{title}](#{anchor})")
        elif f.level == 3 and f.number:
            # Subsection: 1.1.1 概要
            title = f"{f.number} {f.title}"
            anchor = f"{f.number.replace('.', '')}-{f.title.lower().replace(' ', '-')}"
            toc_lines.append(f"    - [{title}](#{anchor})")
        elif not f.number:
            # Special sections like 謝辞, 参考文献
            if any(kw in f.title for kw in ["謝辞", "参考文献", "付録"]):
                anchor = f.title.lower().replace(' ', '-')
                toc_lines.append(f"- [{f.title}](#{anchor})")
    
    return "\n".join(toc_lines)


def convert_image_paths(content: str) -> str:
    """
    Convert image paths from thesis-source format to output format.

    Supported input formats:
    - ![alt](screenshot/prover/default.png)
    - ![alt](../src/assets/screenshot/prover/default.png)
    - ![alt](prover/default.png)  (if file exists in screenshot dir)
    - ![alt](image.png)  (if file exists in assets dir)

    Output format:
    - ![alt](../../../assets/screenshot/prover/default.png)
    - ![alt](../../../assets/image.png)

    Also validates that referenced images exist.
    """
    # Assets directory (parent of screenshot dir)
    assets_dir = SCREENSHOT_DIR.parent

    def replace_path(match: re.Match) -> str:
        alt_text = match.group(1)
        original_path = match.group(2)

        # Skip URLs
        if original_path.startswith(('http://', 'https://', '//')):
            return match.group(0)

        # Normalize path - extract the screenshot-relative portion
        path_lower = original_path.lower()

        # Find screenshot-related path components
        screenshot_components = ['prover', 'verifier', 'registrar-console', 'executive-console']

        for component in screenshot_components:
            if component in path_lower:
                # Extract from component onwards
                idx = original_path.lower().find(component)
                relative_path = original_path[idx:]

                # Verify file exists
                full_path = SCREENSHOT_DIR / relative_path
                if full_path.exists():
                    # Convert to output format
                    # Output is in: src/content/docs/research/thesis-vX-Y.md
                    # Screenshots are in: src/assets/screenshot/
                    # Relative path: ../../../assets/screenshot/
                    new_path = f"../../../assets/screenshot/{relative_path}"
                    return f"![{alt_text}]({new_path})"
                else:
                    print(f"  ⚠️  Warning: Screenshot not found: {full_path}")
                    return match.group(0)

        # If path already contains 'assets/screenshot', adjust it
        if 'assets/screenshot' in original_path:
            idx = original_path.find('assets/screenshot')
            new_path = "../../../" + original_path[idx:]
            return f"![{alt_text}]({new_path})"

        # If path already contains 'assets/', adjust it
        if 'assets/' in original_path:
            idx = original_path.find('assets/')
            new_path = "../../../" + original_path[idx:]
            return f"![{alt_text}]({new_path})"

        # Check if file exists directly in assets directory
        # (e.g., tankyu-chart.png -> ../../../assets/tankyu-chart.png)
        filename = Path(original_path).name
        if (assets_dir / filename).exists():
            new_path = f"../../../assets/{filename}"
            return f"![{alt_text}]({new_path})"

        # Return unchanged if we can't process it
        print(f"  ⚠️  Warning: Image not found: {original_path}")
        return match.group(0)

    return IMAGE_PATH_PATTERN.sub(replace_path, content)


def add_paragraph_indentation(content: str) -> str:
    """
    Add Japanese-style paragraph indentation (全角スペース) to paragraphs.

    This adds a full-width space (　) at the beginning of paragraphs for proper
    Japanese academic paper formatting. Works for both Astro HTML output and
    Word/PDF export.

    Rules:
    - Indent ALL paragraph lines by default
    - Skip only: headings, list items, code blocks, tables, images, HTML tags,
      blockquotes, empty lines, horizontal rules, already indented lines
    """
    lines = content.split('\n')
    result = []
    in_code_block = False
    in_table = False

    def is_list_item(s: str) -> bool:
        """Check if line is a list item (not bold/italic markdown)."""
        # Unordered list: "- ", "* ", "+ " (with space after)
        if s.startswith('- ') or s.startswith('+ '):
            return True
        # Asterisk list item: "* " but not "**" (bold)
        if s.startswith('* ') and not s.startswith('**'):
            return True
        # Ordered list: "1. ", "2. ", etc.
        if re.match(r'^\d+\.\s', s):
            return True
        return False

    for i, line in enumerate(lines):
        # Track code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            result.append(line)
            continue

        if in_code_block:
            result.append(line)
            continue

        # Track tables (lines starting with |)
        stripped = line.strip()
        if stripped.startswith('|'):
            in_table = True
            result.append(line)
            continue
        elif in_table and not stripped.startswith('|') and stripped:
            in_table = False

        if in_table:
            result.append(line)
            continue

        # Skip conditions - everything else gets indented
        skip_conditions = [
            not stripped,                           # Empty line
            stripped.startswith('#'),               # Heading
            is_list_item(stripped),                 # List items (-, * , +, 1.)
            stripped.startswith('>'),               # Blockquote
            stripped.startswith('<'),               # HTML tag
            stripped.startswith('!'),               # Image
            stripped.startswith('['),               # Link at start
            stripped.startswith('　'),              # Already indented
            stripped.startswith('---'),             # Horizontal rule
            stripped.startswith('___'),             # Horizontal rule
        ]

        if any(skip_conditions):
            result.append(line)
            continue

        # Indent all other lines (paragraphs)
        leading_ws = line[:len(line) - len(line.lstrip())]
        result.append(f"{leading_ws}　{stripped}")

    return '\n'.join(result)


def build_thesis_content(files: list[SourceFile], config: configparser.ConfigParser) -> str:
    """Build complete thesis content from source files."""
    frontmatter_parts = []
    body_parts = []
    current_chapter = None
    
    for f in files:
        content = f.content.strip()
        
        # Separate frontmatter from body
        # Frontmatter: level 0 or filename starts with "_"
        if f.level == 0 or f.path.name.startswith("_"):
            if content:
                frontmatter_parts.append(content)
            continue
        
        # Generate heading based on level
        heading = None
        if f.number:
            if f.level == 1:
                # Chapter: # 第1章 序論
                heading = f"# 第{f.number}章 {f.title}"
                current_chapter = f.number
            elif f.level == 2:
                # Section: ## 1.1 研究の背景
                heading = f"## {f.number} {f.title}"
            elif f.level == 3:
                # Subsection: ### 1.1.1 概要
                heading = f"### {f.number} {f.title}"
            else:
                heading = f"#### {f.number} {f.title}"
        elif f.title:
            # Special sections (謝辞, 参考文献, 付録)
            heading = f"# {f.title}"
        
        # Build the output for this file
        if heading:
            if content:
                # Has both heading and content
                if not content.startswith("#"):
                    body_parts.append(f"{heading}\n\n{content}")
                else:
                    # Content already has heading, use as-is
                    body_parts.append(content)
            else:
                # Only heading (directory entry)
                body_parts.append(heading)
        elif content:
            # Only content
            body_parts.append(content)
    
    # Build: frontmatter + TOC + body
    result_parts = []
    
    # Add frontmatter (title, author info)
    if frontmatter_parts:
        result_parts.append("\n\n".join(frontmatter_parts))
        result_parts.append("---")
    
    # Generate and add table of contents
    toc = generate_toc(files)
    result_parts.append(toc)
    result_parts.append("---")
    
    # Add body content with --- only before chapters (# headings)
    # Skip first --- since TOC already ends with ---
    body_output = []
    first_chapter = True
    for part in body_parts:
        # Remove trailing --- from content (will be added explicitly between chapters)
        part = re.sub(r'\n---\s*$', '', part)
        
        # Add --- before chapter headings (# 第X章 or special sections like # 謝辞)
        if part.startswith("# "):
            if not first_chapter:
                body_output.append("---")
            first_chapter = False
        body_output.append(part)
    
    result_parts.append("\n\n".join(body_output))

    # Apply paragraph indentation (全角スペース)
    result = "\n\n".join(result_parts)
    result = add_paragraph_indentation(result)

    return result


# =============================================================================
# Init Command
# =============================================================================

def init_source_structure(config: configparser.ConfigParser) -> None:
    """Initialize the thesis source directory structure."""
    source_dir = get_source_dir(config)
    
    if source_dir.exists():
        print(f"Source directory already exists: {source_dir}")
        response = input("Do you want to reset it? This will delete existing content. [y/N]: ")
        if response.lower() != "y":
            print("Cancelled.")
            return
        shutil.rmtree(source_dir)
    
    print(f"Creating thesis source structure at {source_dir}...")
    
    # Create directory structure
    structure = {
        "_frontmatter.md": """# 修士論文

## 文書の真正性検証フレームワークの設計と実装
### ゼロ知識証明を用いた分散型証明システム Tri-CertFramework

神戸情報大学院大学
情報技術研究科 情報システム専攻

学籍番号：XXXXX
氏名：XXXXX

提出日：20XX年X月X日
""",
        "第1章 序論/1.1 研究の背景.md": """
私たちの社会は、様々な「文書」の信頼性に支えられている。

（ここに内容を記述）
""",
        "第1章 序論/1.2 真正性検証とは.md": """
本研究における「真正性検証」とは、以下の3つの要素を確認するプロセスを指す。

（ここに内容を記述）
""",
        "第2章 先行研究と技術基盤/2.1 デジタル証明書の現状.md": """
現在、デジタル証明書の代表的な規格としてX.509証明書が広く使用されている。

（ここに内容を記述）
""",
        "謝辞.md": """# 謝辞

本研究は、神戸情報大学院大学情報技術研究科情報システム専攻において実施しました。

（ここに謝辞を記述）
""",
        "参考文献.md": """# 参考文献

[1] 参考文献1
[2] 参考文献2
""",
    }
    
    for path_str, content in structure.items():
        file_path = source_dir / path_str
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_text(content.strip() + "\n", encoding="utf-8")
        print(f"  ✓ Created {path_str}")
    
    print(f"\n✅ Initialized thesis source structure")
    print(f"   Edit files in: {source_dir.relative_to(DOCS_DIR)}")
    print(f"   Then run: ./thesis.py build")


# =============================================================================
# Split Command
# =============================================================================

def split_existing_thesis(config: configparser.ConfigParser) -> None:
    """Split an existing thesis markdown into directory structure."""
    revisions = get_revisions()
    
    if not revisions:
        print("Error: No existing thesis revisions found.")
        print("Use './thesis.py init' to create a new structure.")
        sys.exit(1)
    
    # Get the current (latest) revision
    current = next((r for r in revisions if r.is_current), revisions[0])
    
    print(f"Splitting thesis v{current.version} into source structure...")
    
    source_dir = get_source_dir(config)
    if source_dir.exists():
        response = input(f"Source directory exists. Overwrite? [y/N]: ")
        if response.lower() != "y":
            print("Cancelled.")
            return
        shutil.rmtree(source_dir)
    
    source_dir.mkdir(parents=True)
    
    # Read thesis content
    content = current.file_path.read_text(encoding="utf-8")
    
    # Remove frontmatter
    if content.startswith("---"):
        end_match = re.search(r"\n---\n", content[3:])
        if end_match:
            content = content[end_match.end() + 3:].strip()
    
    # Split by chapters (# 第X章 or # Title)
    chapter_pattern = re.compile(r"^# (.+)$", re.MULTILINE)
    section_pattern = re.compile(r"^## (\d+\.\d+)\s+(.+)$", re.MULTILINE)
    subsection_pattern = re.compile(r"^### (\d+\.\d+\.\d+)\s+(.+)$", re.MULTILINE)
    
    # Keywords to skip (not real chapters)
    skip_keywords = ["修士論文", "目次", "論文", "Thesis", "Table of Contents"]
    
    # Find all chapters
    all_chapters = list(chapter_pattern.finditer(content))
    
    # Filter out non-chapter headings
    chapters = []
    for ch in all_chapters:
        title = ch.group(1)
        if not any(kw in title for kw in skip_keywords):
            chapters.append(ch)
    
    for i, chapter_match in enumerate(chapters):
        chapter_title = chapter_match.group(1)
        chapter_start = chapter_match.end()
        chapter_end = chapters[i + 1].start() if i + 1 < len(chapters) else len(content)
        chapter_content = content[chapter_start:chapter_end].strip()
        
        # Determine chapter number and directory name
        num_match = re.match(r"第(\d+)章\s+(.+)", chapter_title)
        if num_match:
            chapter_num = num_match.group(1)
            chapter_name = num_match.group(2)
            dir_name = f"第{chapter_num}章 {chapter_name}"
        else:
            # Special chapters like 謝辞, 参考文献
            if any(kw in chapter_title for kw in ["謝辞", "参考文献", "付録"]):
                file_path = source_dir / f"{chapter_title}.md"
                file_path.write_text(f"# {chapter_title}\n\n{chapter_content}\n", encoding="utf-8")
                print(f"  ✓ Created {chapter_title}.md")
                continue
            dir_name = chapter_title
            chapter_num = "0"
        
        chapter_dir = source_dir / dir_name
        chapter_dir.mkdir(exist_ok=True)
        
        # Find sections within this chapter
        sections = list(section_pattern.finditer(chapter_content))
        
        if sections:
            for j, section_match in enumerate(sections):
                section_num = section_match.group(1)
                section_title = section_match.group(2)
                section_start = section_match.end()
                section_end = sections[j + 1].start() if j + 1 < len(sections) else len(chapter_content)
                section_content = chapter_content[section_start:section_end].strip()
                
                # Check for subsections
                subsections = list(subsection_pattern.finditer(section_content))
                
                if subsections:
                    # Create section directory (sanitize for filesystem)
                    section_dir_name = sanitize_filename(f"{section_num} {section_title}")
                    section_dir = chapter_dir / section_dir_name
                    section_dir.mkdir(exist_ok=True)
                    
                    for k, subsec_match in enumerate(subsections):
                        subsec_num = subsec_match.group(1)
                        subsec_title = subsec_match.group(2)
                        subsec_start = subsec_match.end()
                        subsec_end = subsections[k + 1].start() if k + 1 < len(subsections) else len(section_content)
                        subsec_content = section_content[subsec_start:subsec_end].strip()
                        
                        # Sanitize filename
                        subsec_filename = sanitize_filename(f"{subsec_num} {subsec_title}.md")
                        file_path = section_dir / subsec_filename
                        file_path.write_text(subsec_content + "\n", encoding="utf-8")
                        print(f"  ✓ Created {file_path.relative_to(source_dir)}")
                else:
                    # No subsections, create section file (sanitize for filesystem)
                    section_filename = sanitize_filename(f"{section_num} {section_title}.md")
                    file_path = chapter_dir / section_filename
                    file_path.write_text(section_content + "\n", encoding="utf-8")
                    print(f"  ✓ Created {file_path.relative_to(source_dir)}")
        else:
            # No sections, save whole chapter content
            file_path = chapter_dir / f"_chapter.md"
            file_path.write_text(chapter_content + "\n", encoding="utf-8")
            print(f"  ✓ Created {file_path.relative_to(source_dir)}")
    
    # Extract frontmatter (content before first real chapter)
    original_content = current.file_path.read_text(encoding="utf-8")
    
    # Find first real chapter (# 第X章) by searching for the pattern
    first_real_chapter = re.search(r"^# 第\d+章\s+", original_content, re.MULTILINE)
    
    if first_real_chapter:
        frontmatter_content = original_content[:first_real_chapter.start()].strip()
        
        # Remove YAML frontmatter
        if frontmatter_content.startswith("---"):
            end_match = re.search(r"\n---\n", frontmatter_content[3:])
            if end_match:
                frontmatter_content = frontmatter_content[end_match.end() + 3:].strip()
        
        # Remove :::caution blocks (past version notices)
        frontmatter_content = re.sub(
            r":::caution\[.*?\].*?:::",
            "",
            frontmatter_content,
            flags=re.DOTALL
        ).strip()
        
        # Remove table of contents section (will be auto-generated)
        # TOC is between "# 目次" and the next "---" or end
        toc_match = re.search(r"# 目次\s*\n.*?(?=\n---\s*$|\Z)", frontmatter_content, re.MULTILINE | re.DOTALL)
        if toc_match:
            frontmatter_content = frontmatter_content[:toc_match.start()].strip()
        
        # Clean up any trailing ---
        frontmatter_content = re.sub(r"\n---\s*$", "", frontmatter_content).strip()

        if frontmatter_content:
            fm_file = source_dir / "_frontmatter.md"
            fm_file.write_text(frontmatter_content + "\n", encoding="utf-8")
            print(f"  ✓ Created _frontmatter.md")
    
    print(f"\n✅ Split thesis into source structure")
    print(f"   Directory: {source_dir.relative_to(DOCS_DIR)}")
    print(f"   Run './thesis.py tree' to see the structure")


# =============================================================================
# Build Command
# =============================================================================

def build_thesis(config: configparser.ConfigParser) -> None:
    """Build thesis from source directory."""
    source_dir = get_source_dir(config)
    
    if not source_dir.exists():
        print(f"Error: Source directory not found: {source_dir}")
        print("Run './thesis.py init' to create the structure, or")
        print("Run './thesis.py split' to split an existing thesis.")
        sys.exit(1)
    
    # Collect source files
    files = collect_source_files(source_dir)
    
    if not files:
        print("Error: No source files found.")
        sys.exit(1)
    
    print(f"Building thesis from {len(files)} source files...")
    
    # Build content
    content = build_thesis_content(files, config)
    
    # Increment minor version
    major, minor = get_current_version(config)
    new_minor = minor + 1
    set_version(config, major, new_minor)
    
    new_rev = Revision(
        major=major,
        minor=new_minor,
        file_path=RESEARCH_DIR / f"thesis-v{major}-{new_minor}.md",
        is_current=True,
    )
    
    # Build frontmatter
    frontmatter = f'''---
title: "修士論文 v{new_rev.version}"
description: "{config.get('metadata', 'subtitle')} - {config.get('metadata', 'description')}"
tableOfContents:
  minHeadingLevel: 2
  maxHeadingLevel: 3
---

'''
    
    # Convert mermaid blocks
    content = re.sub(
        r"```mermaid\n(.*?)```",
        r'<pre class="mermaid">\n\1</pre>',
        content,
        flags=re.DOTALL,
    )

    # Convert image paths to correct relative paths
    content = convert_image_paths(content)

    # Write output
    new_rev.file_path.write_text(frontmatter + content, encoding="utf-8")
    print(f"  ✓ Created {new_rev.filename}")
    
    # Mark old revisions as past
    for rev in get_revisions():
        if rev.file_path != new_rev.file_path:
            mark_as_past_version(rev, new_rev)
    
    # Update index.md
    update_index_md_for_new_version(new_rev)
    print(f"  ✓ Updated index.md")
    
    # Update astro.config.mjs
    update_astro_config_for_new_version(new_rev)
    print(f"  ✓ Updated astro.config.mjs")
    
    print(f"\n✅ Built thesis v{new_rev.version}")
    print(f"   File: {new_rev.file_path.relative_to(DOCS_DIR)}")
    print(f"   Size: {new_rev.file_path.stat().st_size / 1024:.1f} KB")


# =============================================================================
# Export Commands (DOCX/PDF)
# =============================================================================

MERMAID_PATTERN = re.compile(r'```mermaid\n(.*?)```', re.DOTALL)
MERMAID_PRE_PATTERN = re.compile(r'<pre class="mermaid">\n?(.*?)</pre>', re.DOTALL)

# JPEG quality for export
# Higher quality for screenshots with text to prevent blurriness
# 95+ is essentially lossless for text, 92-94 is very high quality
JPEG_QUALITY_HIGH = 95  # For screenshots with text (near lossless)
JPEG_QUALITY_DEFAULT = 88  # For diagrams and other images

# Compression threshold: only compress images larger than this width (in pixels)
# Images smaller than this will NOT be resized to avoid upscaling artifacts
# Retina displays typically produce 2x resolution, so 1400px is a safe threshold
COMPRESSION_THRESHOLD_WIDTH = 1400

# Target width for Word export (in pixels)
# Only applied to images exceeding COMPRESSION_THRESHOLD_WIDTH
# Higher resolution for print quality (論文印刷用)
TARGET_WIDTH_SETTINGS = {
    # Screenshots - higher resolution for text clarity in print
    # (Retina screenshots are typically 1800+ px, target 1400px for Word)
    "prover": 1400,
    "verifier": 1400,
    "registrar-console": 1400,
    "executive-console": 1400,
    # Diagrams - keep original size (usually small)
    "mermaid": None,  # None = no resize
    # Charts and figures
    "tankyu-chart": 1200,
    # Default for other images
    "default": 1200,
}


def get_target_width(image_path: Path) -> Optional[int]:
    """
    Determine the target width for an image based on its path/name.

    Returns:
        Target width in pixels, or None if image should not be resized.
    """
    path_str = str(image_path).lower()
    name = image_path.stem.lower()

    # Check for specific image types
    for key, width in TARGET_WIDTH_SETTINGS.items():
        if key == "default":
            continue
        if key in path_str or key in name:
            return width

    return TARGET_WIDTH_SETTINGS["default"]


def should_compress_image(image_path: Path, img_width: int) -> bool:
    """
    Determine if an image should be compressed/resized.

    Only large images (above COMPRESSION_THRESHOLD_WIDTH) are compressed.
    This prevents quality degradation of small images like Mermaid diagrams.

    Args:
        image_path: Path to the image file
        img_width: Current width of the image in pixels

    Returns:
        True if image should be compressed, False otherwise.
    """
    # Don't resize Mermaid diagrams - they're already optimized
    path_str = str(image_path).lower()
    if "mermaid" in path_str:
        return False

    # Only compress if image exceeds threshold
    return img_width > COMPRESSION_THRESHOLD_WIDTH


def get_jpeg_quality(image_path: Path) -> int:
    """
    Determine JPEG quality based on image type.

    Screenshots with text use higher quality to prevent blurriness.
    """
    path_str = str(image_path).lower()
    # Screenshots typically contain text and need higher quality
    screenshot_keywords = ["prover", "verifier", "registrar", "executive", "screenshot"]
    if any(kw in path_str for kw in screenshot_keywords):
        return JPEG_QUALITY_HIGH
    return JPEG_QUALITY_DEFAULT


def convert_png_to_jpg(png_path: Path, output_dir: Path, force_resize: bool = False) -> Path:
    """
    Convert a PNG image to JPEG format for smaller file size.

    Only resizes images that exceed COMPRESSION_THRESHOLD_WIDTH to preserve
    quality of smaller images (like Mermaid diagrams).

    Args:
        png_path: Path to the source PNG file
        output_dir: Directory to save the converted file
        force_resize: If True, always resize regardless of threshold

    Returns the path to the converted JPEG file.
    """
    from PIL import Image

    jpg_name = png_path.stem + ".jpg"
    jpg_path = output_dir / jpg_name

    try:
        with Image.open(png_path) as img:
            original_width = img.width
            original_height = img.height

            # Convert RGBA to RGB (JPEG doesn't support alpha)
            if img.mode in ('RGBA', 'LA', 'P'):
                # Create white background
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')

            # Only resize if image exceeds threshold or force_resize is True
            should_resize = force_resize or should_compress_image(png_path, original_width)

            if should_resize:
                target_width = get_target_width(png_path)
                if target_width is not None and original_width > target_width:
                    # Calculate new height maintaining aspect ratio
                    ratio = target_width / original_width
                    new_height = int(original_height * ratio)
                    img = img.resize((target_width, new_height), Image.Resampling.LANCZOS)

            # Use appropriate JPEG quality
            quality = get_jpeg_quality(png_path)
            img.save(jpg_path, 'JPEG', quality=quality, optimize=True)

        return jpg_path
    except Exception as e:
        print(f"  ⚠️  Warning: Failed to convert {png_path.name} to JPEG: {e}")
        return png_path  # Return original on failure


def check_dependencies() -> dict[str, bool]:
    """Check if required dependencies are installed."""
    deps = {}

    # Check Pillow (for image conversion)
    try:
        from PIL import Image
        deps["pillow"] = True
    except ImportError:
        deps["pillow"] = False

    # Check pandoc
    try:
        result = subprocess.run(["pandoc", "--version"], capture_output=True, text=True)
        deps["pandoc"] = result.returncode == 0
    except FileNotFoundError:
        deps["pandoc"] = False

    # Check mermaid-cli (mmdc)
    try:
        result = subprocess.run(["mmdc", "--version"], capture_output=True, text=True)
        deps["mermaid"] = result.returncode == 0
    except FileNotFoundError:
        deps["mermaid"] = False

    return deps


def prerender_mermaid(content: str, output_dir: Path, use_jpeg: bool = False) -> str:
    """
    Convert Mermaid code blocks to images.

    Handles both:
    - ```mermaid ... ``` (raw markdown)
    - <pre class="mermaid"> ... </pre> (converted format)

    Args:
        content: Markdown content with Mermaid blocks
        output_dir: Directory to save rendered images
        use_jpeg: If True, convert PNG to JPEG for smaller file size
    """
    mermaid_dir = output_dir / "mermaid"
    mermaid_dir.mkdir(parents=True, exist_ok=True)

    counter = [0]  # Use list to allow modification in nested function

    def replace_mermaid(match: re.Match) -> str:
        mermaid_code = match.group(1).strip()
        counter[0] += 1
        png_name = f"mermaid-{counter[0]}.png"
        png_path = mermaid_dir / png_name

        # Write mermaid code to temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
            f.write(mermaid_code)
            mmd_path = f.name

        try:
            # Run mmdc to convert to PNG (scale 1.5 for balance between quality and size)
            result = subprocess.run(
                ["mmdc", "-i", mmd_path, "-o", str(png_path), "-b", "white", "-s", "1.5"],
                capture_output=True,
                text=True,
            )
            if result.returncode != 0:
                print(f"  ⚠️  Warning: Mermaid rendering failed for diagram {counter[0]}")
                print(f"      {result.stderr[:200]}")
                return match.group(0)  # Return original on failure
        finally:
            Path(mmd_path).unlink(missing_ok=True)

        # Convert to JPEG if requested
        if use_jpeg and png_path.exists():
            jpg_path = convert_png_to_jpg(png_path, mermaid_dir)
            png_path.unlink(missing_ok=True)  # Remove original PNG
            return f"![]({jpg_path})"

        # Return markdown image reference (empty alt text to avoid duplication with caption)
        return f"![]({png_path})"

    # Replace both formats
    content = MERMAID_PATTERN.sub(replace_mermaid, content)
    content = MERMAID_PRE_PATTERN.sub(replace_mermaid, content)

    if counter[0] > 0:
        fmt = "JPEG" if use_jpeg else "PNG"
        print(f"  ✓ Rendered {counter[0]} Mermaid diagram(s) as {fmt}")

    return content


def convert_image_paths_for_export(content: str) -> str:
    """
    Convert image paths to absolute paths for Pandoc export.

    Input: ![alt](../../../assets/screenshot/prover/default.png)
    Output: ![alt](/absolute/path/to/screenshot/prover/default.png)
    """
    assets_dir = DOCS_DIR / "src" / "assets"

    def replace_path(match: re.Match) -> str:
        alt_text = match.group(1)
        original_path = match.group(2)

        # Skip URLs and already absolute paths
        if original_path.startswith(('http://', 'https://', '//')):
            return match.group(0)

        # If already absolute, return as-is
        if original_path.startswith('/'):
            return match.group(0)

        # Handle relative paths from thesis output format
        # ../../../assets/screenshot/prover/default.png -> assets/screenshot/prover/default.png
        if 'assets/' in original_path:
            idx = original_path.find('assets/')
            relative_from_assets = original_path[idx + 7:]  # After "assets/"
            absolute_path = assets_dir / relative_from_assets
            if absolute_path.exists():
                return f"![{alt_text}]({absolute_path})"

        # Handle direct screenshot references
        screenshot_components = ['prover', 'verifier', 'registrar-console', 'executive-console']
        for component in screenshot_components:
            if component in original_path.lower():
                idx = original_path.lower().find(component)
                relative_path = original_path[idx:]
                full_path = SCREENSHOT_DIR / relative_path
                if full_path.exists():
                    return f"![{alt_text}]({full_path})"

        # Handle direct filename references (e.g., tankyu-chart.png)
        # Try to find in assets directory directly
        filename = Path(original_path).name
        direct_path = assets_dir / filename
        if direct_path.exists():
            return f"![{alt_text}]({direct_path})"

        # Also try in screenshot directory
        screenshot_path = SCREENSHOT_DIR / filename
        if screenshot_path.exists():
            return f"![{alt_text}]({screenshot_path})"

        print(f"  ⚠️  Warning: Image not found: {original_path}")
        return match.group(0)

    return IMAGE_PATH_PATTERN.sub(replace_path, content)


def convert_images_to_jpeg_for_export(content: str, output_dir: Path) -> str:
    """
    Convert PNG images to JPEG format for smaller file size.

    Only compresses large images (above COMPRESSION_THRESHOLD_WIDTH) to preserve
    quality of smaller images. Screenshots with text use higher JPEG quality.

    Processes all image references in the content, converts PNG files to JPEG,
    and updates the paths in the content.
    """
    from PIL import Image

    assets_dir = DOCS_DIR / "src" / "assets"
    jpeg_dir = output_dir / "images"
    jpeg_dir.mkdir(parents=True, exist_ok=True)

    # Statistics
    stats = {"converted": 0, "resized": 0, "skipped_small": 0, "reused": 0}

    def replace_and_convert(match: re.Match) -> str:
        alt_text = match.group(1)
        original_path = match.group(2)

        # Skip URLs
        if original_path.startswith(('http://', 'https://', '//')):
            return match.group(0)

        # Skip already converted JPEG/JPG files
        if original_path.lower().endswith(('.jpg', '.jpeg')):
            return match.group(0)

        # Find the actual file path
        actual_path = None

        # If already absolute path
        if original_path.startswith('/'):
            actual_path = Path(original_path)
        # Handle relative paths from thesis output format
        elif 'assets/' in original_path:
            idx = original_path.find('assets/')
            relative_from_assets = original_path[idx + 7:]
            actual_path = assets_dir / relative_from_assets
        else:
            # Handle direct screenshot references
            screenshot_components = ['prover', 'verifier', 'registrar-console', 'executive-console']
            for component in screenshot_components:
                if component in original_path.lower():
                    idx = original_path.lower().find(component)
                    relative_path = original_path[idx:]
                    actual_path = SCREENSHOT_DIR / relative_path
                    break

            if actual_path is None or not actual_path.exists():
                # Try direct filename in assets
                filename = Path(original_path).name
                if (assets_dir / filename).exists():
                    actual_path = assets_dir / filename
                elif (SCREENSHOT_DIR / filename).exists():
                    actual_path = SCREENSHOT_DIR / filename

        # If file not found or not PNG, return original
        if actual_path is None or not actual_path.exists():
            return match.group(0)

        if not actual_path.suffix.lower() == '.png':
            # Non-PNG file, just return with absolute path
            return f"![{alt_text}]({actual_path})"

        # Check image size to determine if compression is needed
        try:
            with Image.open(actual_path) as img:
                original_width = img.width
        except Exception:
            original_width = 0

        # Use unique name based on file path hash
        unique_name = actual_path.stem + "_" + str(hash(str(actual_path)) % 10000)
        jpg_path = jpeg_dir / (unique_name + ".jpg")

        if jpg_path.exists():
            # Already converted in a previous run
            stats["reused"] += 1
            return f"![{alt_text}]({jpg_path})"

        # Convert PNG to JPEG
        jpg_path = convert_png_to_jpg(actual_path, jpeg_dir)

        if jpg_path != actual_path:  # Conversion succeeded
            # Rename to unique name if needed
            if jpg_path.name != unique_name + ".jpg":
                new_jpg_path = jpeg_dir / (unique_name + ".jpg")
                if jpg_path.exists():
                    jpg_path.rename(new_jpg_path)
                    jpg_path = new_jpg_path

            stats["converted"] += 1

            # Check if image was resized
            if should_compress_image(actual_path, original_width):
                stats["resized"] += 1
            else:
                stats["skipped_small"] += 1

        return f"![{alt_text}]({jpg_path})"

    result = IMAGE_PATH_PATTERN.sub(replace_and_convert, content)

    # Print statistics
    total = stats["converted"] + stats["reused"]
    if total > 0:
        print(f"  ✓ Processed {total} image(s):")
        if stats["converted"] > 0:
            print(f"      - Converted: {stats['converted']} ({stats['resized']} resized, {stats['skipped_small']} kept original size)")
        if stats["reused"] > 0:
            print(f"      - Reused from cache: {stats['reused']}")

    return result


def build_export_content(files: list[SourceFile], config: configparser.ConfigParser) -> str:
    """
    Build thesis content for export (DOCX/PDF).

    Similar to build_thesis_content but:
    - No frontmatter YAML
    - No mermaid <pre> conversion (handled separately)
    - Different separator handling
    """
    body_parts = []

    for f in files:
        content = f.content.strip()

        # Skip frontmatter files for body (will be handled separately)
        if f.level == 0 or f.path.name.startswith("_"):
            continue

        # Generate heading based on level
        heading = None
        if f.number:
            if f.level == 1:
                heading = f"# 第{f.number}章 {f.title}"
            elif f.level == 2:
                heading = f"## {f.number} {f.title}"
            elif f.level == 3:
                heading = f"### {f.number} {f.title}"
            else:
                heading = f"#### {f.number} {f.title}"
        elif f.title:
            if any(kw in f.title for kw in ["謝辞", "参考文献", "付録"]):
                heading = f"# {f.title}"

        if heading:
            if content:
                if not content.startswith("#"):
                    body_parts.append(f"{heading}\n\n{content}")
                else:
                    body_parts.append(content)
            else:
                body_parts.append(heading)
        elif content:
            body_parts.append(content)

    # Apply paragraph indentation (全角スペース)
    result = "\n\n".join(body_parts)
    result = add_paragraph_indentation(result)

    return result


def get_frontmatter_content(files: list[SourceFile]) -> str:
    """Extract frontmatter content from source files."""
    for f in files:
        if f.path.name == "_frontmatter.md":
            return f.content.strip()
    return ""


def export_docx(config: configparser.ConfigParser) -> Optional[Path]:
    """Export thesis to DOCX format."""
    deps = check_dependencies()
    if not deps["pandoc"]:
        print("Error: Pandoc is not installed.")
        print("  Install with: brew install pandoc")
        return None

    source_dir = get_source_dir(config)
    if not source_dir.exists():
        print(f"Error: Source directory not found: {source_dir}")
        return None

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("Exporting thesis to DOCX...")

    # Collect and build content
    files = collect_source_files(source_dir)
    content = build_export_content(files, config)
    frontmatter = get_frontmatter_content(files)

    # Pre-render Mermaid diagrams if mermaid-cli is available
    if deps["mermaid"]:
        content = prerender_mermaid(content, OUTPUT_DIR)
    else:
        print("  ⚠️  Warning: mermaid-cli not installed, Mermaid diagrams will be skipped")
        # Remove mermaid blocks
        content = MERMAID_PATTERN.sub('[Mermaid diagram - install mmdc to render]', content)
        content = MERMAID_PRE_PATTERN.sub('[Mermaid diagram - install mmdc to render]', content)

    # Convert image paths to absolute
    content = convert_image_paths_for_export(content)

    # Combine frontmatter and body
    full_content = f"{frontmatter}\n\n---\n\n{content}"

    # Write to temp file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(full_content)
        temp_md = f.name

    # Get version for filename
    major, minor = get_current_version(config)
    output_file = OUTPUT_DIR / f"thesis-v{major}-{minor}.docx"

    try:
        # Build pandoc command
        cmd = [
            "pandoc",
            temp_md,
            "--from", "markdown",
            "--to", "docx",
            "--toc",
            "--toc-depth=3",
            "-o", str(output_file),
        ]

        # Add reference doc if available
        if REFERENCE_DOCX.exists():
            cmd.extend(["--reference-doc", str(REFERENCE_DOCX)])

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            print(f"Error: Pandoc failed")
            print(result.stderr)
            return None

        print(f"  ✓ Created {output_file.name}")
        print(f"    Size: {output_file.stat().st_size / 1024:.1f} KB")
        return output_file

    finally:
        Path(temp_md).unlink(missing_ok=True)


def export_pdf(config: configparser.ConfigParser) -> Optional[Path]:
    """Export thesis to PDF format using XeLaTeX."""
    deps = check_dependencies()
    if not deps["pandoc"]:
        print("Error: Pandoc is not installed.")
        print("  Install with: brew install pandoc")
        return None

    # Check for XeLaTeX
    try:
        result = subprocess.run(["xelatex", "--version"], capture_output=True, text=True)
        has_xelatex = result.returncode == 0
    except FileNotFoundError:
        has_xelatex = False

    if not has_xelatex:
        print("Error: XeLaTeX is not installed.")
        print("  Install with: brew install --cask mactex-no-gui")
        return None

    source_dir = get_source_dir(config)
    if not source_dir.exists():
        print(f"Error: Source directory not found: {source_dir}")
        return None

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("Exporting thesis to PDF...")

    # Collect and build content
    files = collect_source_files(source_dir)
    content = build_export_content(files, config)
    frontmatter = get_frontmatter_content(files)

    # Pre-render Mermaid diagrams
    if deps["mermaid"]:
        content = prerender_mermaid(content, OUTPUT_DIR)
    else:
        print("  ⚠️  Warning: mermaid-cli not installed, Mermaid diagrams will be skipped")
        content = MERMAID_PATTERN.sub('[Mermaid diagram - install mmdc to render]', content)
        content = MERMAID_PRE_PATTERN.sub('[Mermaid diagram - install mmdc to render]', content)

    # Convert image paths to absolute
    content = convert_image_paths_for_export(content)

    # Combine frontmatter and body
    full_content = f"{frontmatter}\n\n---\n\n{content}"

    # Write to temp file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(full_content)
        temp_md = f.name

    # Get version for filename
    major, minor = get_current_version(config)
    output_file = OUTPUT_DIR / f"thesis-v{major}-{minor}.pdf"

    try:
        # Build pandoc command for PDF
        # Use XeTeX with native Unicode support (works with basictex)
        # Japanese fonts are handled directly by XeTeX without xeCJK
        cmd = [
            "pandoc",
            temp_md,
            "--from", "markdown",
            "--to", "pdf",
            "--pdf-engine=xelatex",
            "-V", "documentclass=article",
            "-V", "geometry:top=2.5cm,left=2.5cm,right=2cm,bottom=2cm,a4paper",
            "-V", "mainfont=Hiragino Mincho ProN",
            "-V", "sansfont=Hiragino Sans",
            "-V", "monofont=Menlo",
            "--toc",
            "--toc-depth=3",
            "-o", str(output_file),
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            print(f"Error: Pandoc failed")
            print(result.stderr[:500])
            return None

        print(f"  ✓ Created {output_file.name}")
        print(f"    Size: {output_file.stat().st_size / 1024:.1f} KB")
        return output_file

    finally:
        Path(temp_md).unlink(missing_ok=True)


def export_thesis(config: configparser.ConfigParser) -> None:
    """Export thesis to both Word and PDF formats."""
    print("\n📄 Thesis Export")
    print("=" * 50)

    word_path = build_word(config)
    print()
    pdf_path = export_pdf(config)

    print("\n" + "-" * 50)
    if word_path:
        print(f"✅ Word: {word_path.relative_to(SCRIPT_DIR)}")
    else:
        print("❌ Word: Failed")

    if pdf_path:
        print(f"✅ PDF:  {pdf_path.relative_to(SCRIPT_DIR)}")
    else:
        print("❌ PDF:  Failed")


# =============================================================================
# DOCX XML Post-Processing
# =============================================================================

def postprocess_docx_xml(docx_path: Path) -> Path:
    """
    Post-process DOCX file at XML level to fix page break issues.

    This function uses string replacement instead of XML parsing to avoid
    corrupting the DOCX file structure. Changes made:
    1. Fix keepNext: change w:val="0" to enabled (remove val attribute)
    2. Remove lastRenderedPageBreak markers (rendering hints that cause issues)
    3. Fix widowControl: change w:val="0" to enabled

    Returns the path to the processed DOCX file.
    """
    import zipfile
    import re

    print("  → Post-processing DOCX at XML level (string replacement)...")

    # Create temp directory for extraction
    temp_dir = docx_path.parent / '_docx_temp'
    temp_dir.mkdir(exist_ok=True)

    try:
        # Extract DOCX
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        # Read document.xml as text
        doc_xml_path = temp_dir / 'word' / 'document.xml'
        with open(doc_xml_path, 'r', encoding='utf-8') as f:
            content = f.read()

        original_content = content
        keepnext_fixed = 0
        pagebreak_removed = 0
        widowcontrol_fixed = 0

        # 1. Fix keepNext: change <w:keepNext w:val="0"/> to <w:keepNext/>
        # This enables keepNext for all paragraphs where it was explicitly disabled
        keepnext_patterns = [
            r'<w:keepNext\s+w:val="0"\s*/>',
            r'<w:keepNext\s+w:val="false"\s*/>',
            r'<w:keepNext w:val="0"/>',
            r'<w:keepNext w:val="false"/>',
        ]
        for pattern in keepnext_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            keepnext_fixed += len(matches)
            content = re.sub(pattern, '<w:keepNext/>', content, flags=re.IGNORECASE)

        # 2. Remove lastRenderedPageBreak markers
        # These are just rendering hints from previous Word sessions and can cause issues
        pagebreak_patterns = [
            r'<w:lastRenderedPageBreak\s*/>',
            r'<w:lastRenderedPageBreak/>',
        ]
        for pattern in pagebreak_patterns:
            matches = re.findall(pattern, content)
            pagebreak_removed += len(matches)
            content = re.sub(pattern, '', content)

        # 3. Fix widowControl: change <w:widowControl w:val="0"/> to <w:widowControl/>
        widowcontrol_patterns = [
            r'<w:widowControl\s+w:val="0"\s*/>',
            r'<w:widowControl\s+w:val="false"\s*/>',
            r'<w:widowControl w:val="0"/>',
            r'<w:widowControl w:val="false"/>',
        ]
        for pattern in widowcontrol_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            widowcontrol_fixed += len(matches)
            content = re.sub(pattern, '<w:widowControl/>', content, flags=re.IGNORECASE)

        # 4. Remove empty paragraphs that might cause layout issues
        # These can appear as <w:p><w:pPr/></w:p> or <w:p></w:p> with no content
        empty_para_removed = 0

        # 5. Fix section breaks that cause editing issues
        # Remove continuous section breaks that aren't needed (can cause DEL key issues)
        # Pattern: <w:sectPr><w:type w:val="continuous"/></w:sectPr> or similar
        section_break_fixed = 0
        # Look for empty sectPr elements or ones with only type="continuous"
        empty_sectpr_patterns = [
            r'<w:sectPr\s*>\s*<w:type\s+w:val="continuous"\s*/>\s*</w:sectPr>',
            r'<w:sectPr><w:type w:val="continuous"/></w:sectPr>',
        ]
        for pattern in empty_sectpr_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            section_break_fixed += len(matches)
            content = re.sub(pattern, '', content, flags=re.IGNORECASE)

        # 6. Remove orphaned page breaks within runs that can cause issues
        # Pattern: <w:br w:type="page"/> in places where it shouldn't be
        orphan_break_fixed = 0

        # Report changes
        if keepnext_fixed > 0:
            print(f"    ✓ Fixed keepNext for {keepnext_fixed} element(s)")
        if pagebreak_removed > 0:
            print(f"    ✓ Removed {pagebreak_removed} lastRenderedPageBreak marker(s)")
        if widowcontrol_fixed > 0:
            print(f"    ✓ Fixed widowControl for {widowcontrol_fixed} element(s)")
        if section_break_fixed > 0:
            print(f"    ✓ Removed {section_break_fixed} unnecessary section break(s)")

        # Only write if changes were made
        if content != original_content:
            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(content)

            # Repackage DOCX preserving original ZIP structure
            import os
            new_docx_path = docx_path.parent / f'{docx_path.stem}_processed{docx_path.suffix}'

            with zipfile.ZipFile(new_docx_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root_dir, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = Path(root_dir) / file
                        arcname = file_path.relative_to(temp_dir)
                        zipf.write(file_path, arcname)

            # Replace original with processed version
            docx_path.unlink()
            new_docx_path.rename(docx_path)
            print("    ✓ DOCX file updated successfully")
        else:
            print("    ℹ️  No changes needed")

        return docx_path

    finally:
        # Clean up temp directory
        import shutil
        if temp_dir.exists():
            shutil.rmtree(temp_dir)


# =============================================================================
# Word Export with Cover Template
# =============================================================================

def build_word(config: configparser.ConfigParser) -> Optional[Path]:
    """
    Build Word document with cover template.

    This creates a properly formatted Word document with:
    - Cover page from template (表紙.docx)
    - Table of contents
    - Body content from thesis source
    - Proper margins and page numbering

    Uses docxcompose for proper document merging (preserves images, styles).
    """
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docxcompose.composer import Composer

    deps = check_dependencies()
    if not deps["pandoc"]:
        print("Error: Pandoc is not installed.")
        print("  Install with: brew install pandoc")
        return None

    if not deps["pillow"]:
        print("Error: Pillow is not installed (required for image compression).")
        print("  Install with: pip install Pillow")
        return None

    source_dir = get_source_dir(config)
    if not source_dir.exists():
        print(f"Error: Source directory not found: {source_dir}")
        return None

    if not COVER_DOCX.exists():
        print(f"Error: Cover template not found: {COVER_DOCX}")
        return None

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("📝 Building Word document with cover template...")

    # Collect and build content
    files = collect_source_files(source_dir)
    content = build_export_content(files, config)

    # Pre-render Mermaid diagrams if mermaid-cli is available
    # Use JPEG format for smaller file size
    if deps["mermaid"]:
        # Keep Mermaid diagrams as PNG to avoid conversion artifacts
        # JPEG conversion can cause cropping/scaling issues with diagrams
        content = prerender_mermaid(content, OUTPUT_DIR, use_jpeg=False)
    else:
        print("  ⚠️  Warning: mermaid-cli not installed, Mermaid diagrams will be skipped")
        content = MERMAID_PATTERN.sub('[Mermaid diagram - install mmdc to render]', content)
        content = MERMAID_PRE_PATTERN.sub('[Mermaid diagram - install mmdc to render]', content)

    # Convert all PNG images to JPEG for smaller file size
    print("  → Converting images to JPEG...")
    content = convert_images_to_jpeg_for_export(content, OUTPUT_DIR)

    # Process figure/table captions for Word export
    # Convert <p align="center"><strong>図X.X: ...</strong></p> to proper Pandoc format
    # This ensures captions are properly formatted in Word
    def process_captions(text: str) -> str:
        import re
        # Pattern for HTML captions: <p align="center"><strong>図/表X.X: ...</strong></p>
        caption_pattern = re.compile(
            r'<p\s+align="center">\s*<strong>((?:図|表)\d+\.\d+:\s*[^<]+)</strong>\s*</p>',
            re.IGNORECASE
        )
        # Replace with Pandoc-friendly centered text
        # Using a div with custom style that will be processed later
        text = caption_pattern.sub(r'\n\n::: {.figure-caption}\n**\1**\n:::\n\n', text)
        return text

    # Remove alt text from images that have a caption following them
    # This must be done BEFORE process_captions to match the original HTML format
    # Pattern: ![alt text](path)\n\n<p align="center"><strong>図X.X: caption</strong></p>
    # Replace with: ![](path)\n\n<p align="center"><strong>図X.X: caption</strong></p>
    def remove_duplicate_alt_text(text: str) -> str:
        # Pattern: image with alt text followed by figure caption (with possible whitespace)
        pattern = re.compile(
            r'!\[([^\]]+)\]\(([^)]+)\)(\s*\n\s*\n?\s*<p\s+align="center">\s*<strong>(?:図|表)\d+\.\d+:)',
            re.IGNORECASE
        )
        # Replace with empty alt text
        text = pattern.sub(r'![](\2)\3', text)
        return text

    content = remove_duplicate_alt_text(content)
    content = process_captions(content)

    # Remove horizontal rules (---) as they may not be appropriate for academic papers
    content = re.sub(r'\n---\n', '\n\n', content)
    content = re.sub(r'^---\n', '\n', content)
    content = re.sub(r'\n---$', '\n', content)

    # Add spacing between chapters and sections for readability
    # Add blank lines before chapter headings (# 第X章)
    content = re.sub(r'\n(# 第\d+章)', r'\n\n\n\1', content)
    # Add blank lines before section headings (## X.X)
    content = re.sub(r'\n(## \d+\.\d+)', r'\n\n\1', content)
    # Add blank lines before subsection headings (### X.X.X)
    content = re.sub(r'\n(### \d+\.\d+\.\d+)', r'\n\n\1', content)
    # Add blank lines before special chapters (謝辞, 参考文献, 付録)
    content = re.sub(r'\n(# (?:謝辞|参考文献|付録))', r'\n\n\n\1', content)

    # Write body content to temp file (without frontmatter - cover handles that)
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(content)
        temp_md = f.name

    # Get version for filename
    major, minor = get_current_version(config)
    body_temp = OUTPUT_DIR / "_body_temp.docx"
    cover_temp = OUTPUT_DIR / "_cover_temp.docx"
    output_file = OUTPUT_DIR / f"thesis-v{major}-{minor}.docx"

    try:
        # Step 1: Convert body content to DOCX with Pandoc
        print("  → Converting body content with Pandoc...")
        cmd = [
            "pandoc",
            temp_md,
            "--from", "markdown",
            "--to", "docx",
            "--toc",
            "--toc-depth=3",
            "-o", str(body_temp),
        ]

        # Use reference doc for styling if available
        if REFERENCE_DOCX.exists():
            cmd.extend(["--reference-doc", str(REFERENCE_DOCX)])

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            print(f"Error: Pandoc failed")
            print(result.stderr)
            return None

        print("  ✓ Body content converted")

        # Step 2: Prepare cover document with page break
        print("  → Preparing cover page...")
        cover_doc = Document(COVER_DOCX)

        # Apply margins (per hard_format.md)
        for section in cover_doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

        # Add page break after cover content
        cover_doc.add_page_break()
        cover_doc.save(cover_temp)
        print("  ✓ Cover page prepared")

        # Step 3: Merge documents using docxcompose
        print("  → Merging documents...")
        master = Document(cover_temp)
        composer = Composer(master)
        body_doc = Document(body_temp)
        composer.append(body_doc)

        # Step 4: Save merged document temporarily
        merged_temp = OUTPUT_DIR / "_merged_temp.docx"
        composer.save(merged_temp)

        # Step 5: Post-process the merged document
        print("  → Post-processing document...")
        final_doc = Document(merged_temp)

        # Helper function to add section break with page numbering
        def add_section_break_with_page_restart(para):
            """Add section break before paragraph and restart page numbering."""
            sectPr = OxmlElement('w:sectPr')
            # Page number format - restart from 1
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(qn('w:start'), '1')
            sectPr.append(pgNumType)
            para._p.addprevious(sectPr)

        # Find chapter headings and add page breaks
        chapter1_para = None
        chapter1_idx = -1
        chapter_paragraphs = []  # List of (paragraph, chapter_name) for chapters after 第1章

        for i, para in enumerate(final_doc.paragraphs):
            text = para.text.strip()

            # Find 第1章
            if text.startswith('第1章') and chapter1_para is None:
                chapter1_para = para
                chapter1_idx = i
                print(f"    ✓ Found 第1章 at paragraph {i}")

            # Find other chapter headings (第2章, 第3章, etc.) and special sections
            # Must be a heading style OR short text (chapter titles are short, body text is long)
            # This prevents matching body text like "第2章で分析した既存手法の..."
            elif re.match(r'^第[2-9]\d*章\s+\S', text):
                style_name = para.style.name if para.style else ''
                is_heading_style = 'Heading' in style_name or '見出し' in style_name
                is_short = len(text) < 50  # Chapter titles are typically short
                if is_heading_style or is_short:
                    chapter_paragraphs.append((para, text))
            elif text in ['謝辞', '参考文献', '付録']:
                chapter_paragraphs.append((para, text))

            # Center-align figure/table captions
            if text.startswith('図') or text.startswith('表'):
                # Check if it looks like a caption (図X.X: or 表X.X:)
                if re.match(r'^[図表]\d+\.\d+:', text):
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Center-align paragraphs with figure-caption style or containing only bold caption
            if para.style and 'caption' in para.style.name.lower():
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Insert section break before 第1章 for page numbering
        if chapter1_para:
            # Add section break (next page) before 第1章
            sectPr = OxmlElement('w:sectPr')
            sectType = OxmlElement('w:type')
            sectType.set(qn('w:val'), 'nextPage')
            sectPr.append(sectType)
            # Page number restart from 1
            pgNumType = OxmlElement('w:pgNumType')
            pgNumType.set(qn('w:start'), '1')
            sectPr.append(pgNumType)
            # Insert before the paragraph
            chapter1_para._p.addprevious(sectPr)
            print("    ✓ Added section break before 第1章 with page restart from 1")

        # Add page breaks before other chapters (第2章 onwards and special sections)
        for chapter_para, chapter_name in chapter_paragraphs:
            # Add page break before this chapter using paragraph property
            pPr = chapter_para._p.get_or_add_pPr()
            # Remove existing pageBreakBefore if present
            for existing in list(pPr):
                if existing.tag.endswith('pageBreakBefore'):
                    pPr.remove(existing)
            # Add pageBreakBefore element
            pageBreakBefore = OxmlElement('w:pageBreakBefore')
            pPr.insert(0, pageBreakBefore)

        if chapter_paragraphs:
            chapter_names = [name for _, name in chapter_paragraphs]
            print(f"    ✓ Added page breaks before {len(chapter_paragraphs)} chapter(s): {', '.join(chapter_names[:3])}{'...' if len(chapter_names) > 3 else ''}")

        # Process images: center them and limit maximum width
        # Maximum width for images in Word (in EMUs - English Metric Units)
        # A4 paper with 2.5cm left and 2cm right margins = 16cm usable width
        # 1 cm = 360000 EMUs, so max width = 16 * 360000 = 5,760,000 EMUs
        MAX_IMAGE_WIDTH_EMU = 5_760_000
        # For smaller images (like Mermaid diagrams), don't stretch them
        # Mermaid diagrams are typically ~700px, which at 96dpi = ~7.3 inches = ~18.5cm
        # But in Word, we want them smaller, so cap at ~12cm for readability
        MAX_SMALL_IMAGE_WIDTH_EMU = 4_320_000  # 12cm

        image_resize_count = 0
        for para in final_doc.paragraphs:
            # Check if paragraph contains only an image (drawing)
            drawings = para._element.findall('.//' + qn('w:drawing'))
            if drawings and not para.text.strip():
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Limit image width if too large
                for drawing in drawings:
                    # Find extent element (defines displayed size)
                    extents = drawing.findall('.//' + qn('wp:extent'))
                    for extent in extents:
                        cx = int(extent.get('cx', 0))
                        cy = int(extent.get('cy', 0))

                        if cx > MAX_IMAGE_WIDTH_EMU:
                            # Scale down proportionally
                            ratio = MAX_IMAGE_WIDTH_EMU / cx
                            new_cx = int(cx * ratio)
                            new_cy = int(cy * ratio)
                            extent.set('cx', str(new_cx))
                            extent.set('cy', str(new_cy))
                            image_resize_count += 1

                            # Also update inline extent if present
                            inline_extents = drawing.findall('.//' + qn('a:ext'))
                            for inline_ext in inline_extents:
                                inline_ext.set('cx', str(new_cx))
                                inline_ext.set('cy', str(new_cy))

        if image_resize_count > 0:
            print(f"    ✓ Resized {image_resize_count} oversized image(s) to fit page width")

        # Center-align tables
        for table in final_doc.tables:
            # Set table alignment to center
            tbl = table._tbl
            tblPr = tbl.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Remove existing jc element if present
            existing_jc = tblPr.find(qn('w:jc'))
            if existing_jc is not None:
                tblPr.remove(existing_jc)

            # Add center alignment
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tblPr.append(jc)

        print(f"    ✓ Centered {len(final_doc.tables)} table(s)")

        # Fix table layout issues: prevent tables from jumping to next page
        # when there's enough space, and allow row splitting if needed
        for table in final_doc.tables:
            tblPr = table._tbl.tblPr
            if tblPr is not None:
                # Remove cantSplit at table level if exists
                for child in list(tblPr):
                    if child.tag.endswith('cantSplit'):
                        tblPr.remove(child)

            # For each row, allow splitting across pages
            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                # Remove existing cantSplit
                for child in list(trPr):
                    if child.tag.endswith('cantSplit'):
                        trPr.remove(child)
                # Add cantSplit=false to allow row to split across pages
                cantSplit = OxmlElement('w:cantSplit')
                cantSplit.set(qn('w:val'), '0')
                trPr.append(cantSplit)

                # Also disable page break before for cells
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.page_break_before = False
                        para.paragraph_format.keep_with_next = False

        print(f"    ✓ Fixed table row splitting for {len(final_doc.tables)} table(s)")

        # Fix paragraphs immediately BEFORE tables (table captions)
        # Use Document Body traversal to find paragraphs directly preceding tables
        # and set keepNext at XML level to ensure caption stays with table
        before_table_para_count = 0
        body = final_doc._body._body
        children = list(body)
        for i, child in enumerate(children):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tbl' and i > 0:
                # Check the previous element
                prev_child = children[i - 1]
                prev_tag = prev_child.tag.split('}')[-1] if '}' in prev_child.tag else prev_child.tag
                if prev_tag == 'p':
                    # This paragraph is immediately before a table
                    # Set keepNext at XML level to force Word to keep them together
                    pPr = prev_child.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        prev_child.insert(0, pPr)

                    # Remove existing keepNext and add new one
                    for existing in list(pPr):
                        if existing.tag.endswith('keepNext'):
                            pPr.remove(existing)
                    keepNext = OxmlElement('w:keepNext')
                    # No w:val attribute means true (default)
                    pPr.insert(0, keepNext)

                    # Also remove pageBreakBefore if present
                    for existing in list(pPr):
                        if existing.tag.endswith('pageBreakBefore'):
                            pPr.remove(existing)

                    # Disable widow control for this paragraph
                    for existing in list(pPr):
                        if existing.tag.endswith('widowControl'):
                            pPr.remove(existing)
                    widowControl = OxmlElement('w:widowControl')
                    widowControl.set(qn('w:val'), '0')
                    pPr.append(widowControl)

                    before_table_para_count += 1

        if before_table_para_count > 0:
            print(f"    ✓ Fixed keepNext for {before_table_para_count} paragraph(s) before tables")

        # Fix paragraphs immediately after tables to prevent unnecessary page breaks
        # This addresses the issue where Word pushes content to the next page even when
        # there's enough space after a table
        after_table_para_count = 0
        body = final_doc._body._body
        prev_was_table = False
        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tbl':
                prev_was_table = True
            elif tag == 'p' and prev_was_table:
                # This is a paragraph immediately after a table
                # Find or create pPr element
                pPr = child.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    child.insert(0, pPr)

                # Remove existing widowControl if present and set to false
                for existing in list(pPr):
                    if existing.tag.endswith('widowControl'):
                        pPr.remove(existing)
                widowControl = OxmlElement('w:widowControl')
                widowControl.set(qn('w:val'), '0')
                pPr.append(widowControl)

                # Remove existing keepNext if present (we don't want to keep with next)
                for existing in list(pPr):
                    if existing.tag.endswith('keepNext'):
                        pPr.remove(existing)

                # Remove existing keepLines if present
                for existing in list(pPr):
                    if existing.tag.endswith('keepLines'):
                        pPr.remove(existing)

                after_table_para_count += 1
                prev_was_table = False
            else:
                prev_was_table = False

        if after_table_para_count > 0:
            print(f"    ✓ Fixed widow/orphan control for {after_table_para_count} paragraph(s) after tables")

        # Set code block line spacing to fixed 9-10pt for compactness
        # Code blocks typically use styles like "Source Code", "Code", or have monospace fonts
        from docx.shared import Twips
        code_para_count = 0
        code_para_indices = []  # Track code paragraph indices for later table wrapping
        for i, para in enumerate(final_doc.paragraphs):
            style_name = para.style.name if para.style else ''
            # Check if it's a code-related style
            if any(code_style in style_name.lower() for code_style in ['code', 'source', 'verbatim', 'literal']):
                # Set line spacing to fixed 10pt (200 twips = 10pt, 1pt = 20 twips)
                para.paragraph_format.line_spacing = Pt(10)
                para.paragraph_format.line_spacing_rule = 4  # WD_LINE_SPACING.EXACTLY = 4
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                # Keep with next to prevent page breaks within code blocks
                para.paragraph_format.keep_with_next = True
                para.paragraph_format.keep_together = True
                code_para_count += 1
                code_para_indices.append(i)

        if code_para_count > 0:
            print(f"    ✓ Set line spacing for {code_para_count} code block paragraph(s)")

        # Add borders to code blocks to create a visual box
        # Find consecutive code paragraphs and add borders
        def add_paragraph_borders(para, top=False, bottom=False, left=True, right=True):
            """Add borders to a paragraph using XML."""
            pPr = para._p.get_or_add_pPr()
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is None:
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)

            border_attrs = {
                qn('w:val'): 'single',
                qn('w:sz'): '4',  # 0.5pt
                qn('w:space'): '1',
                qn('w:color'): '808080',  # Gray
            }

            if left:
                left_el = OxmlElement('w:left')
                for k, v in border_attrs.items():
                    left_el.set(k, v)
                pBdr.append(left_el)

            if right:
                right_el = OxmlElement('w:right')
                for k, v in border_attrs.items():
                    right_el.set(k, v)
                pBdr.append(right_el)

            if top:
                top_el = OxmlElement('w:top')
                for k, v in border_attrs.items():
                    top_el.set(k, v)
                pBdr.append(top_el)

            if bottom:
                bottom_el = OxmlElement('w:bottom')
                for k, v in border_attrs.items():
                    bottom_el.set(k, v)
                pBdr.append(bottom_el)

        # Group consecutive code paragraphs and add borders
        if code_para_indices:
            # Find groups of consecutive code paragraphs
            groups = []
            current_group = [code_para_indices[0]]
            for i in range(1, len(code_para_indices)):
                if code_para_indices[i] == code_para_indices[i-1] + 1:
                    current_group.append(code_para_indices[i])
                else:
                    groups.append(current_group)
                    current_group = [code_para_indices[i]]
            groups.append(current_group)

            # Add borders to each group
            for group in groups:
                for i, para_idx in enumerate(group):
                    para = final_doc.paragraphs[para_idx]
                    is_first = (i == 0)
                    is_last = (i == len(group) - 1)
                    add_paragraph_borders(para, top=is_first, bottom=is_last, left=True, right=True)
                    # Add small padding via indentation
                    para.paragraph_format.left_indent = Pt(6)
                    para.paragraph_format.right_indent = Pt(6)

            print(f"    ✓ Added borders to {len(groups)} code block group(s)")

        # Add spacing before section and subsection headings
        # Also disable "page break before" to prevent unwanted page breaks
        section_spacing_count = 0
        heading4_fixed_count = 0

        # First pass: fix Heading 4 style issues
        # Pandoc may not correctly style #### headings, so we detect them by text patterns
        # and apply Heading 4 style manually
        heading4_style = None
        for style in final_doc.styles:
            if style.name == 'Heading 4':
                heading4_style = style
                break

        for para in final_doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else ''

            # Check for section headings (## X.X format -> Heading 2)
            if 'Heading 2' in style_name or re.match(r'^\d+\.\d+\s+\S', text):
                para.paragraph_format.space_before = Pt(18)  # ~1.5 lines before
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.page_break_before = False  # Disable auto page break
                section_spacing_count += 1

            # Check for subsection headings (### X.X.X format -> Heading 3)
            elif 'Heading 3' in style_name or re.match(r'^\d+\.\d+\.\d+\s+\S', text):
                para.paragraph_format.space_before = Pt(12)  # ~1 line before
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.page_break_before = False  # Disable auto page break
                section_spacing_count += 1

            # Check for Heading 4 (#### in source)
            # These might not have Heading 4 style applied by Pandoc
            # Detect by checking if it's a short paragraph that looks like a heading
            elif 'Heading 4' in style_name:
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.page_break_before = False
                section_spacing_count += 1

            # Also check for paragraphs that should be Heading 4 but aren't styled
            # These are typically short lines (< 50 chars) that end without punctuation
            # and are followed by normal paragraphs
            elif (not any(x in style_name.lower() for x in ['heading', 'code', 'caption', 'toc'])
                  and len(text) < 60
                  and len(text) > 5
                  and not text.endswith(('。', '．', '.', '、', '，', ',', '：', ':'))
                  and not text.startswith(('　', '-', '・', '1', '2', '3', '4', '5', '6', '7', '8', '9', '図', '表', '['))
                  and not re.match(r'^第\d+章', text)
                  and not re.match(r'^\d+\.\d+', text)):
                # Check if this looks like a heading (e.g., ends with question or contains "とは" etc.)
                heading_patterns = [
                    r'とは$',
                    r'について$',
                    r'の(概要|特徴|目的|背景|理由|仕組み|方法|種類|比較|課題|利点|欠点)$',
                    r'(SHA|ZKP|FIDO|WebAuthn|Poseidon|Groth16|VKNFT)',  # Technical terms
                    r'が.*理由$',
                    r'は.*か$',  # Question pattern
                    r'の使い分け$',
                ]
                is_likely_heading = any(re.search(p, text) for p in heading_patterns)

                if is_likely_heading and heading4_style:
                    para.style = heading4_style
                    para.paragraph_format.space_before = Pt(10)
                    para.paragraph_format.space_after = Pt(4)
                    para.paragraph_format.page_break_before = False
                    heading4_fixed_count += 1

        if section_spacing_count > 0:
            print(f"    ✓ Adjusted spacing for {section_spacing_count} section heading(s)")
        if heading4_fixed_count > 0:
            print(f"    ✓ Fixed Heading 4 style for {heading4_fixed_count} paragraph(s)")

        # Unify line spacing for all normal paragraphs to prevent inconsistent spacing
        # This addresses the issue where line spacing varies between paragraphs
        normal_para_count = 0
        for para in final_doc.paragraphs:
            style_name = para.style.name if para.style else ''
            text = para.text.strip()

            # Skip headings, code blocks, captions, and empty paragraphs
            if any(skip in style_name.lower() for skip in ['heading', 'code', 'source', 'verbatim', 'literal', 'caption', 'toc']):
                continue
            if re.match(r'^(?:第\d+章|謝辞|参考文献|付録|\d+\.\d+)', text):
                continue
            if not text:
                continue

            # Set consistent line spacing for normal paragraphs
            # Use 1.5 line spacing (approximately 18pt for 12pt font)
            # This creates a balanced, readable layout
            para.paragraph_format.line_spacing = 1.5  # 1.5 lines
            para.paragraph_format.line_spacing_rule = 1  # WD_LINE_SPACING.ONE_AND_A_HALF = 1
            normal_para_count += 1

        if normal_para_count > 0:
            print(f"    ✓ Unified line spacing for {normal_para_count} normal paragraph(s)")

        # Keep list introduction paragraphs with their lists
        # Paragraphs ending with "：" or ":" followed by a list should stay together
        keep_with_list_count = 0
        for para in final_doc.paragraphs:
            text = para.text.strip()
            # Check if paragraph ends with colon (Japanese or English)
            if text.endswith('：') or text.endswith(':'):
                para.paragraph_format.keep_with_next = True
                keep_with_list_count += 1

        if keep_with_list_count > 0:
            print(f"    ✓ Set keep-with-next for {keep_with_list_count} list introduction paragraph(s)")

        # Fix bullet points - remove the "・" prefix from list items
        # This happens when Pandoc converts lists but the style adds Japanese bullets
        for para in final_doc.paragraphs:
            text = para.text
            # Check for patterns like "・ 1." or "・ -" at the start
            if text.startswith('・ '):
                # Get the text after "・ "
                rest = text[2:]
                # Clear all runs and rewrite
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = rest
                else:
                    para.add_run(rest)
            # Also handle "・" without space
            elif text.startswith('・'):
                rest = text[1:]
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = rest
                else:
                    para.add_run(rest)

        # Reduce list indentation from default to minimal
        # List items have numPr in their pPr element
        # Also fix tab stops to reduce space between number and text
        list_indent_count = 0
        for para in final_doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    # This is a list item - reduce indentation
                    ind = pPr.find(qn('w:ind'))
                    if ind is None:
                        ind = OxmlElement('w:ind')
                        pPr.append(ind)
                    # Set minimal left indent (360 twips = ~18pt = about 2 half-width spaces)
                    # Hanging indent determines where text starts after the number
                    ind.set(qn('w:left'), '360')
                    ind.set(qn('w:hanging'), '360')

                    # Remove existing tabs and set a tight tab stop
                    existing_tabs = pPr.find(qn('w:tabs'))
                    if existing_tabs is not None:
                        pPr.remove(existing_tabs)
                    tabs = OxmlElement('w:tabs')
                    tab = OxmlElement('w:tab')
                    tab.set(qn('w:val'), 'left')
                    tab.set(qn('w:pos'), '360')  # Tab at same position as indent
                    tabs.append(tab)
                    pPr.append(tabs)

                    list_indent_count += 1

        if list_indent_count > 0:
            print(f"    ✓ Reduced indentation for {list_indent_count} list item(s)")

        # Keep figure/table captions with their content (prevent page breaks)
        keep_with_count = 0
        for i, para in enumerate(final_doc.paragraphs):
            text = para.text.strip()
            # Figure/table captions should stay with the image above
            if re.match(r'^(?:図|表)\d+\.\d+:', text):
                para.paragraph_format.keep_with_next = False  # Caption doesn't need to keep with next
                # But the paragraph BEFORE the caption (the image) should keep with next
                if i > 0:
                    prev_para = final_doc.paragraphs[i-1]
                    prev_para.paragraph_format.keep_with_next = True
                    keep_with_count += 1

            # Also keep images with their captions
            has_drawing = para._element.findall('.//' + qn('w:drawing'))
            if has_drawing:
                para.paragraph_format.keep_with_next = True
                keep_with_count += 1

        if keep_with_count > 0:
            print(f"    ✓ Set keep-with-next for {keep_with_count} figure/image paragraph(s)")

        # Step 6: Configure sections for page numbering
        # Cover page and TOC should have no page numbers
        # Chapter 1 onwards should have page numbers starting from 1
        sections = list(final_doc.sections)

        if len(sections) >= 1:
            # First section (cover + TOC): no page numbers
            first_section = sections[0]
            footer = first_section.footer
            footer.is_linked_to_previous = False
            # Clear footer content for first section
            for para in footer.paragraphs:
                para.clear()
            print("    ✓ Cleared page numbers from first section (cover + TOC)")

        if len(sections) >= 2:
            # Second section (chapter 1 onwards): add page numbers
            second_section = sections[1]
            footer = second_section.footer
            footer.is_linked_to_previous = False

            # Add page number to footer
            if not footer.paragraphs:
                para = footer.add_paragraph()
            else:
                para = footer.paragraphs[0]
                para.clear()

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add PAGE field with proper structure: begin -> instrText -> separate -> end
            run = para.add_run()

            # Field begin
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar_begin)

            # Instruction text
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = " PAGE "
            run._r.append(instrText)

            # Field separate (required for field to display)
            fldChar_separate = OxmlElement('w:fldChar')
            fldChar_separate.set(qn('w:fldCharType'), 'separate')
            run._r.append(fldChar_separate)

            # Field end
            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar_end)

            print("    ✓ Added page numbers to second section (from 第1章)")

        # Save final document
        final_doc.save(output_file)
        merged_temp.unlink(missing_ok=True)

        # Post-process the DOCX at XML level to fix page break issues
        output_file = postprocess_docx_xml(output_file)

        print(f"  ✓ Created {output_file.name}")
        print(f"    Size: {output_file.stat().st_size / 1024:.1f} KB")

        return output_file

    finally:
        Path(temp_md).unlink(missing_ok=True)
        body_temp.unlink(missing_ok=True)
        cover_temp.unlink(missing_ok=True)


# =============================================================================
# Major Version Command
# =============================================================================

def increment_major_version(config: configparser.ConfigParser) -> None:
    """Increment major version (e.g., 1.5 → 2.0)."""
    major, minor = get_current_version(config)
    new_major = major + 1
    
    print(f"Incrementing major version: {major}.{minor} → {new_major}.0")
    
    response = input("This will create a new major version. Continue? [y/N]: ")
    if response.lower() != "y":
        print("Cancelled.")
        return
    
    set_version(config, new_major, 0)
    print(f"\n✅ Set version to {new_major}.0")
    print("   Run './thesis.py build' to create the new revision.")


# =============================================================================
# Fetch Command
# =============================================================================

def fetch_version(config: configparser.ConfigParser) -> None:
    """Sync thesis.config version with the latest revision in research directory."""
    revisions = get_revisions()
    
    if not revisions:
        print("No thesis revisions found in research directory.")
        print("Config version unchanged.")
        return
    
    # Find the current (latest) revision
    current = next((r for r in revisions if r.is_current), None)
    if not current:
        # If no current, use the highest version
        current = max(revisions, key=lambda r: (r.major, r.minor))
    
    config_major, config_minor = get_current_version(config)
    
    print(f"Research directory:")
    print(f"  Latest revision: v{current.major}.{current.minor}")
    print(f"  Total revisions: {len(revisions)}")
    print(f"\nConfig file:")
    print(f"  Current version: v{config_major}.{config_minor}")
    
    if current.major == config_major and current.minor == config_minor:
        print("\n✅ Config is already in sync.")
        return
    
    print(f"\nSync config to v{current.major}.{current.minor}? [y/N]: ", end="")
    response = input()
    if response.lower() != "y":
        print("Cancelled.")
        return
    
    set_version(config, current.major, current.minor)
    print(f"\n✅ Config synced to v{current.major}.{current.minor}")


# =============================================================================
# Tree Command
# =============================================================================

def show_source_tree(config: configparser.ConfigParser) -> None:
    """Show the thesis source directory structure."""
    source_dir = get_source_dir(config)
    
    if not source_dir.exists():
        print(f"Source directory not found: {source_dir}")
        print("Run './thesis.py init' to create the structure.")
        return
    
    print(f"\n📁 Thesis Source Structure")
    print(f"   {source_dir.relative_to(DOCS_DIR)}/")
    print("=" * 60)
    
    def print_tree(path: Path, prefix: str = "") -> None:
        items = sorted(path.iterdir(), key=lambda p: (
            0 if p.name.startswith("_") else 1,
            not p.is_dir(),
            p.name,
        ))
        
        for i, item in enumerate(items):
            if item.name.startswith("."):
                continue
            
            is_last = i == len(items) - 1
            connector = "└── " if is_last else "├── "
            
            if item.is_dir():
                print(f"{prefix}{connector}📁 {item.name}/")
                new_prefix = prefix + ("    " if is_last else "│   ")
                print_tree(item, new_prefix)
            else:
                size = item.stat().st_size
                size_str = f"({size / 1024:.1f} KB)" if size > 1024 else f"({size} B)"
                print(f"{prefix}{connector}📄 {item.name} {size_str}")
    
    print_tree(source_dir)
    
    # Show stats
    files = collect_source_files(source_dir)
    print("\n" + "-" * 60)
    print(f"Total: {len(files)} source file(s)")


# =============================================================================
# Helper Functions for Revisions
# =============================================================================

def get_revisions() -> list[Revision]:
    """Get all thesis revisions sorted by version (descending)."""
    revisions = []
    
    if not RESEARCH_DIR.exists():
        return revisions
    
    for file_path in RESEARCH_DIR.glob("thesis-v*.md"):
        match = THESIS_FILE_PATTERN.match(file_path.name)
        if match:
            major, minor = int(match.group(1)), int(match.group(2))
            content = file_path.read_text(encoding="utf-8")
            
            is_current = "過去版" not in content[:500]
            
            title_match = re.search(r'title:\s*["\']?([^"\'\n]+)', content[:500])
            title = title_match.group(1) if title_match else ""
            
            revisions.append(Revision(
                major=major,
                minor=minor,
                file_path=file_path,
                is_current=is_current,
                title=title,
            ))
    
    revisions.sort(reverse=True)
    return revisions


def mark_as_past_version(rev: Revision, new_current: Revision) -> None:
    """Mark a revision as past version."""
    content = rev.file_path.read_text(encoding="utf-8")
    
    # Check if already marked
    if "過去版" in content[:500]:
        # Update link to new current
        content = re.sub(
            r"\[v[\d.]+\]\(\./thesis-v\d+-\d+/\)",
            f"[v{new_current.version}](./{new_current.slug}/)",
            content,
        )
        rev.file_path.write_text(content, encoding="utf-8")
        return
    
    # Add badge to frontmatter
    if content.startswith("---"):
        end_match = re.search(r"\n---\n", content[3:])
        if end_match:
            fm_end = end_match.end() + 3
            fm = content[4:end_match.start() + 3]
            body = content[fm_end:]
            
            # Add sidebar badge
            if "sidebar:" not in fm:
                fm += '\nsidebar:\n  badge:\n    text: "過去版"\n    variant: "caution"'
            
            # Add warning
            warning = f"""
:::caution[過去バージョン]
これは過去バージョン（v{rev.version}）です。最新版は [v{new_current.version}](./{new_current.slug}/) をご覧ください。
:::

"""
            if ":::caution[過去バージョン]" not in body:
                body = warning + body.lstrip()
            
            content = f"---\n{fm}\n---\n{body}"
            rev.file_path.write_text(content, encoding="utf-8")


def update_index_md_for_new_version(new_rev: Revision) -> None:
    """Update index.md with new version."""
    if not INDEX_FILE.exists():
        return

    content = INDEX_FILE.read_text(encoding="utf-8")

    # Check if this version already exists
    if f"thesis-v{new_rev.major}-{new_rev.minor}" in content:
        # Already exists, just update status
        content = re.sub(
            rf"\| v{new_rev.version} \| [^|]+ \| [^|]+ \| \[論文を読む →\]\(\./thesis-v{new_rev.major}-{new_rev.minor}/\) \|",
            f"| **v{new_rev.version}** | 最新 | 📗 現行版 | [論文を読む →](./{new_rev.slug}/) |",
            content,
        )
    else:
        # Update existing "最新" row to past version
        content = re.sub(
            r"\| \*\*v([\d.]+)\*\* \| 最新 \| 📗 現行版 \| \[論文を読む →\]\(\./thesis-v(\d+)-(\d+)/\) \|",
            r"| v\1 | - | 📕 過去版 | [論文を読む →](./thesis-v\2-\3/) |",
            content,
        )

        # Add new row
        new_row = f"| **v{new_rev.version}** | 最新 | 📗 現行版 | [論文を読む →](./{new_rev.slug}/) |"
        table_header = "|-----------|--------|------|--------|"
        content = content.replace(table_header, table_header + "\n" + new_row)

    INDEX_FILE.write_text(content, encoding="utf-8")


def update_astro_config_for_new_version(new_rev: Revision) -> None:
    """Update astro.config.mjs with new version."""
    if not ASTRO_CONFIG.exists():
        return

    content = ASTRO_CONFIG.read_text(encoding="utf-8")

    # Check if this version already exists
    if f"research/{new_rev.slug}" in content:
        # Already exists, skip
        return

    # Add new item after '論文一覧'
    new_item = f"            {{ label: '修士論文 v{new_rev.version}', slug: 'research/{new_rev.slug}', translations: {{ en: 'Thesis v{new_rev.version}' }} }},"

    pattern = r"(\{ label: '論文一覧'.*?\},)"
    match = re.search(pattern, content)

    if match:
        insert_pos = match.end()
        content = content[:insert_pos] + "\n" + new_item + content[insert_pos:]
        ASTRO_CONFIG.write_text(content, encoding="utf-8")


# =============================================================================
# List/Show/Remove Commands
# =============================================================================

def list_revisions() -> None:
    """List all thesis revisions."""
    revisions = get_revisions()
    
    if not revisions:
        print("No thesis revisions found.")
        return
    
    config = load_config()
    major, minor = get_current_version(config)
    
    print(f"\n📚 Thesis Revisions (config version: {major}.{minor})")
    print("=" * 70)
    print(f"{'Version':<10} {'Status':<12} {'Filename':<20} {'Size':<10}")
    print("-" * 70)
    
    for rev in revisions:
        status = "📗 Current" if rev.is_current else "📕 Past"
        size = rev.file_path.stat().st_size
        size_str = f"{size / 1024:.1f} KB"
        print(f"v{rev.version:<9} {status:<12} {rev.filename:<20} {size_str:<10}")
    
    print("-" * 70)
    print(f"Total: {len(revisions)} revision(s)")
    print()


def show_revision(version_str: str) -> None:
    """Show details of a specific revision."""
    try:
        major, minor = parse_version_string(version_str)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

    revisions = get_revisions()
    rev = next((r for r in revisions if r.major == major and r.minor == minor), None)
    
    if not rev:
        print(f"Error: Revision v{major}.{minor} not found.")
        sys.exit(1)
    
    content = rev.file_path.read_text(encoding="utf-8")
    
    h1_count = len(re.findall(r"^# ", content, re.MULTILINE))
    h2_count = len(re.findall(r"^## ", content, re.MULTILINE))
    mermaid_count = content.count('<pre class="mermaid">')
    image_count = len(IMAGE_PATH_PATTERN.findall(content))

    print(f"\n📄 Thesis v{rev.version}")
    print("=" * 50)
    print(f"File:       {rev.file_path.relative_to(DOCS_DIR)}")
    print(f"Status:     {'📗 Current' if rev.is_current else '📕 Past'}")
    print(f"Size:       {rev.file_path.stat().st_size / 1024:.1f} KB")
    print(f"Lines:      {len(content.splitlines())}")
    print(f"Chapters:   {h1_count}")
    print(f"Sections:   {h2_count}")
    print(f"Mermaid:    {mermaid_count} diagram(s)")
    print(f"Images:     {image_count} image(s)")
    print()


def sanitize_filename(name: str) -> str:
    """Sanitize a string for use as a filename."""
    # Replace characters that are problematic in filenames
    replacements = {
        "/": "・",
        "\\": "・",
        ":": "：",
        "*": "＊",
        "?": "？",
        '"': "'",
        "<": "＜",
        ">": "＞",
        "|": "｜",
    }
    for old, new in replacements.items():
        name = name.replace(old, new)
    return name


def parse_version_string(version_str: str) -> tuple[int, int]:
    """Parse version from various formats: 1.2, v1.2, v1-2, thesis-v1-2.md, etc."""
    # Remove common prefixes and suffixes
    cleaned = version_str.strip()
    cleaned = re.sub(r"^thesis-", "", cleaned)  # Remove "thesis-" prefix
    cleaned = re.sub(r"\.md$", "", cleaned)      # Remove ".md" suffix
    
    # Try different patterns
    # Format: v1-2 or 1-2
    match = re.match(r"v?(\d+)-(\d+)", cleaned)
    if match:
        return int(match.group(1)), int(match.group(2))
    
    # Format: v1.2 or 1.2
    match = re.match(r"v?(\d+)\.(\d+)", cleaned)
    if match:
        return int(match.group(1)), int(match.group(2))
    
    raise ValueError(f"Invalid version format: {version_str}")


def remove_revision(version_str: str) -> None:
    """Remove a specific revision."""
    try:
        major, minor = parse_version_string(version_str)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
    
    slug = f"thesis-v{major}-{minor}"
    file_path = RESEARCH_DIR / f"{slug}.md"
    
    if not file_path.exists():
        print(f"Error: Revision v{major}.{minor} not found.")
        sys.exit(1)
    
    revisions = get_revisions()
    target = next((r for r in revisions if r.major == major and r.minor == minor), None)
    
    if target and target.is_current:
        print(f"Warning: v{major}.{minor} is the current version!")
        response = input("Are you sure? [y/N]: ")
        if response.lower() != "y":
            print("Cancelled.")
            return
    
    print(f"Removing revision v{major}.{minor}...")
    
    file_path.unlink()
    print(f"  ✓ Deleted {file_path.name}")
    
    # Update index.md
    if INDEX_FILE.exists():
        content = INDEX_FILE.read_text(encoding="utf-8")
        content = re.sub(rf"\|.*?v{major}\.{minor}.*?\|\n", "", content)
        INDEX_FILE.write_text(content, encoding="utf-8")
        print(f"  ✓ Updated index.md")
    
    # Update astro.config.mjs
    if ASTRO_CONFIG.exists():
        content = ASTRO_CONFIG.read_text(encoding="utf-8")
        content = re.sub(rf".*?'修士論文 v{major}\.{minor}'.*?,\n", "", content)
        ASTRO_CONFIG.write_text(content, encoding="utf-8")
        print(f"  ✓ Updated astro.config.mjs")
    
    print(f"\n✅ Removed revision v{major}.{minor}")


# =============================================================================
# Main
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Thesis revision management tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s init              Initialize source directory structure
  %(prog)s split             Split existing thesis into source structure
  %(prog)s build             Build thesis from source (creates new revision)
  %(prog)s major             Increment major version (1.x → 2.0)
  %(prog)s fetch             Sync config with research directory
  %(prog)s tree              Show source directory structure
  %(prog)s list              List all revisions
  %(prog)s show 1.2          Show details of v1.2
  %(prog)s remove 1.0        Remove revision v1.0
        """,
    )
    
    subparsers = parser.add_subparsers(dest="command", help="Available commands")
    
    subparsers.add_parser("init", help="Initialize source directory structure")
    subparsers.add_parser("split", help="Split existing thesis into source structure")
    subparsers.add_parser("build", help="Build thesis from source")
    subparsers.add_parser("major", help="Increment major version")
    subparsers.add_parser("fetch", help="Sync config with research directory")
    subparsers.add_parser("tree", help="Show source directory structure")
    subparsers.add_parser("list", help="List all revisions")
    subparsers.add_parser("pdf", help="Export thesis to PDF format")
    subparsers.add_parser("export", help="Export thesis to both Word and PDF")
    subparsers.add_parser("word", help="Build Word document with cover template")
    
    show_parser = subparsers.add_parser("show", help="Show revision details")
    show_parser.add_argument("version", help="Version (e.g., 1.1)")
    
    remove_parser = subparsers.add_parser("remove", help="Remove a revision")
    remove_parser.add_argument("version", help="Version (e.g., 1.0)")
    
    # Legacy command
    new_parser = subparsers.add_parser("new", help="(Legacy) Create new revision")
    new_parser.add_argument("revision", nargs="?", default="revision")
    
    args = parser.parse_args()
    
    if sys.version_info < (3, 8):
        print("Error: Python 3.8+ required.")
        sys.exit(1)
    
    config = load_config()
    
    if args.command == "init":
        init_source_structure(config)
    elif args.command == "split":
        split_existing_thesis(config)
    elif args.command == "build":
        build_thesis(config)
    elif args.command == "major":
        increment_major_version(config)
    elif args.command == "fetch":
        fetch_version(config)
    elif args.command == "tree":
        show_source_tree(config)
    elif args.command == "list":
        list_revisions()
    elif args.command == "show":
        show_revision(args.version)
    elif args.command == "remove":
        remove_revision(args.version)
    elif args.command == "pdf":
        export_pdf(config)
    elif args.command == "export":
        export_thesis(config)
    elif args.command == "word":
        build_word(config)
    elif args.command == "new":
        # Legacy: just run build
        build_thesis(config)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
