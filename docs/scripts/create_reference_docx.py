#!/usr/bin/env python3
"""
Create a reference DOCX template for Pandoc with university-specified format.

Margins:
- Top: 2.5 cm
- Left: 2.5 cm
- Right: 2 cm
- Bottom: 2 cm

Footer: Page number centered
Fonts: Hiragino Mincho (body), Hiragino Sans (headings)
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

SCRIPT_DIR = Path(__file__).parent
TEMPLATE_PATH = SCRIPT_DIR / "templates" / "reference.docx"


def create_reference_docx():
    """Create reference DOCX with university format specifications."""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

        # Add page number to footer
        footer = section.footer
        footer.is_linked_to_previous = False

        # Create paragraph for page number
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    # Configure styles
    styles = doc.styles

    # Normal style (body text)
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Hiragino Mincho ProN'
    normal_font.size = Pt(10.5)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Mincho ProN')

    # Heading styles
    heading_fonts = ['Hiragino Sans', 'Hiragino Kaku Gothic ProN']

    for i in range(1, 5):
        style_name = f'Heading {i}'
        if style_name in styles:
            heading_style = styles[style_name]
            heading_font = heading_style.font
            heading_font.name = heading_fonts[0]
            heading_font.bold = True
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), heading_fonts[0])

            # Set heading sizes
            if i == 1:
                heading_font.size = Pt(16)
            elif i == 2:
                heading_font.size = Pt(14)
            elif i == 3:
                heading_font.size = Pt(12)
            else:
                heading_font.size = Pt(11)

    # Add sample content for reference (will be replaced by actual content)
    doc.add_heading('見出し1のサンプル', level=1)
    doc.add_paragraph('本文のサンプルテキストです。このテンプレートは大学指定のフォーマットに準拠しています。')

    doc.add_heading('見出し2のサンプル', level=2)
    doc.add_paragraph('段落のサンプルです。')

    doc.add_heading('見出し3のサンプル', level=3)
    doc.add_paragraph('小見出しの下の本文です。')

    # Save
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TEMPLATE_PATH)
    print(f"✓ Created reference template: {TEMPLATE_PATH}")


if __name__ == "__main__":
    create_reference_docx()
